import json
from docx import Document

document = Document()

# read json file
with open("data/厌恶及其他_output_origin.json", "r", encoding="utf-8") as f:
    data = json.load(f)
for item in data:
    page = item["page"]
    content = item["content"]
    # document.add_heading(f"Page {page}", level=1)
    document.add_paragraph(content)
document.save("data/厌恶及其他_output.docx")